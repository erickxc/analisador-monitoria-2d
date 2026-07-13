"""
Exportação dos relatórios em PDF e Word, como alternativa ao Excel.

Diferença importante e proposital: PDF e Word são formatos de LEITURA (um
relatório para imprimir, assinar ou enviar por e-mail), não de manipulação de
dados. Por isso cada tabela é limitada a um número máximo de linhas — a base
completa, com milhares de linhas (ex.: segmentação ABC de 7 mil clientes),
só faz sentido consultar em Excel. O relatório em PDF/Word avisa quando uma
tabela foi cortada.
"""

import os
from datetime import datetime

from reportlab.pdfbase.pdfmetrics import stringWidth

from recursos import CAMINHO_LOGO, NOME_SISTEMA, NOME_EMPRESA

MAX_LINHAS_TABELA = 50

COR_MARCA = "#1F4E78"
COR_MARCA_CLARA = "#EAF0F6"
COR_TEXTO_SECUNDARIO = "#6B7280"
COR_LINHA_SUTIL = "#E2E5E9"
COR_CABECALHO_TABELA = "#000000"
COR_LINHA_ALTERNADA = "#F2F2F2"


def _limitar(df):
    if df is None or df.empty:
        return df, 0
    total = len(df)
    if total > MAX_LINHAS_TABELA:
        return df.head(MAX_LINHAS_TABELA), total
    return df, total


def _formatar_moeda_br(valor):
    try:
        numero = float(valor)
    except (TypeError, ValueError):
        return str(valor)
    texto = f"{numero:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {texto}"


def _formatar_numero_br(valor):
    texto = f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return texto


def _titulo_coluna(nome):
    """
    Nome de coluna pronto pra exibir num cabeçalho de tabela. A maioria das
    colunas já vem com nome legível ("Receita Atual"), mas algumas usam
    PascalCase_Com_Underscore internamente (ex.: "Poder_De_Compra",
    "Percentual_Acumulado") e apareciam assim mesmo, cru, no PDF/Word —
    só o título da seção passava por esse tratamento, não o cabeçalho da
    tabela em si.
    """
    return str(nome).replace("_", " ")


def _largura_maior_palavra(texto, fonte="Helvetica-Bold", tamanho=8.5):
    """
    Largura (em pontos) da maior palavra de um texto, na fonte/tamanho dados.
    Usado para garantir que a coluna nunca fique mais estreita que a maior
    palavra do próprio cabeçalho — sem isso, um cabeçalho longo numa coluna
    numérica estreita ("Períodos Consecutivos em Queda") quebra no meio da
    palavra ("Co" / "nsecutivos") em vez de só trocar de linha no espaço.
    """
    palavras = str(texto).split()
    if not palavras:
        return 0
    return max(stringWidth(palavra, fonte, tamanho) for palavra in palavras)


def _largura_tipica_conteudo(valores, fonte="Helvetica", tamanho=8.5):
    """
    Largura (pontos) representativa do conteúdo de uma coluna de TEXTO — não
    a maior palavra isolada (isso é só a garantia contra quebra no meio da
    palavra do cabeçalho) nem o valor mais longo de todos (uma única razão
    social gigante dominaria a largura da coluna inteira, empurrando as
    outras) — o percentil 90 dos comprimentos reais da coluna. Usado só pra
    colunas de texto; colunas numéricas usam o valor MÁXIMO mesmo (ver
    chamador) porque números não têm outliers de comprimento do mesmo jeito
    que nomes de cliente/produto, e um número cortando no meio fica pior que
    uma coluna de texto um pouco mais larga que o necessário.
    """
    textos = [str(v) for v in valores if v is not None and str(v) != "nan"]
    if not textos:
        return 0
    larguras = sorted(stringWidth(texto, fonte, tamanho) for texto in textos)
    indice = min(int(len(larguras) * 0.9), len(larguras) - 1)
    return larguras[indice]


def exportar_relatorio_pdf(caminho_saida, resultados_analise, nomes_analise, nome_usuario="", colunas_moeda_por_analise=None, nome_empresa="", descricao_analise=None):
    import pandas as pd
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, PageBreak,
        HRFlowable, KeepTogether,
    )
    from reportlab.platypus.tableofcontents import TableOfContents

    colunas_moeda_por_analise = colunas_moeda_por_analise or {}
    descricao_analise = descricao_analise or {}
    cor_marca = colors.HexColor(COR_MARCA)
    cor_marca_clara = colors.HexColor(COR_MARCA_CLARA)
    cor_texto_secundario = colors.HexColor(COR_TEXTO_SECUNDARIO)
    cor_linha_sutil = colors.HexColor(COR_LINHA_SUTIL)
    cor_cabecalho_tabela = colors.HexColor(COR_CABECALHO_TABELA)
    cor_linha_alternada = colors.HexColor(COR_LINHA_ALTERNADA)

    largura_pagina, altura_pagina = landscape(A4)
    margem_lateral = 1.6 * cm
    largura_util = largura_pagina - 2 * margem_lateral

    def _cabecalho_rodape(canvas, _doc):
        canvas.saveState()
        # cabeçalho: nome do sistema discreto no topo direito
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(cor_texto_secundario)
        canvas.drawRightString(largura_pagina - margem_lateral, altura_pagina - 1.15 * cm, NOME_SISTEMA)
        canvas.setStrokeColor(cor_marca_clara)
        canvas.setLineWidth(0.8)
        canvas.line(margem_lateral, altura_pagina - 1.35 * cm, largura_pagina - margem_lateral, altura_pagina - 1.35 * cm)

        # rodapé: empresa + numeração de página
        canvas.setStrokeColor(cor_linha_sutil)
        canvas.setLineWidth(0.6)
        canvas.line(margem_lateral, 1.25 * cm, largura_pagina - margem_lateral, 1.25 * cm)
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(cor_texto_secundario)
        canvas.drawString(margem_lateral, 0.85 * cm, f"{NOME_EMPRESA} — documento de uso interno")
        canvas.drawRightString(largura_pagina - margem_lateral, 0.85 * cm, f"Página {canvas.getPageNumber() - 1}")
        canvas.restoreState()

    def _texto_indice_de(flowable):
        """
        Procura o marcador _indice_texto (ver mais abaixo) no flowable, ou
        num dos flowables dentro dele se for um KeepTogether — o título de
        cada seção vem embrulhado num KeepTogether junto da granularidade/
        descrição/separador (pra nunca ficar "órfão" no fim de uma página),
        e afterFlowable só recebe o KeepTogether inteiro, não cada item dele.
        """
        marcador = getattr(flowable, "_indice_texto", None)
        if marcador:
            return marcador
        for item in getattr(flowable, "_content", None) or []:
            marcador = getattr(item, "_indice_texto", None)
            if marcador:
                return marcador
        return None

    class _DocumentoComIndice(SimpleDocTemplate):
        def afterFlowable(self, flowable):
            texto_indice = _texto_indice_de(flowable)
            if texto_indice:
                self.notify("TOCEntry", (0, texto_indice, self.page))

    doc = _DocumentoComIndice(
        caminho_saida, pagesize=landscape(A4),
        topMargin=2.1 * cm, bottomMargin=1.8 * cm,
        leftMargin=margem_lateral, rightMargin=margem_lateral,
    )

    estilos = getSampleStyleSheet()
    estilo_empresa_capa = ParagraphStyle(
        "EmpresaCapa", parent=estilos["Normal"], fontName="Helvetica-Bold",
        fontSize=16, textColor=cor_marca, alignment=TA_CENTER, spaceBefore=10,
    )
    estilo_meta_capa = ParagraphStyle(
        "MetaCapa", parent=estilos["Normal"], fontName="Helvetica",
        fontSize=9.5, textColor=cor_texto_secundario, alignment=TA_CENTER, spaceBefore=3,
    )
    estilo_rotulo_metadados = ParagraphStyle(
        "RotuloMetadados", parent=estilos["Normal"], fontName="Helvetica",
        fontSize=9, textColor=cor_texto_secundario, alignment=TA_CENTER,
    )
    estilo_valor_metadados = ParagraphStyle(
        "ValorMetadados", parent=estilos["Normal"], fontName="Helvetica-Bold",
        fontSize=13, textColor=cor_marca, alignment=TA_CENTER,
    )
    estilo_secao = ParagraphStyle(
        "TituloSecao", parent=estilos["Heading2"], fontName="Helvetica-Bold",
        fontSize=13, textColor=cor_marca, spaceBefore=0, spaceAfter=6, leading=16,
    )
    estilo_granularidade = ParagraphStyle(
        "Granularidade", parent=estilos["Normal"], fontName="Helvetica",
        fontSize=9.5, textColor=cor_texto_secundario,
    )
    estilo_aviso = ParagraphStyle(
        "Aviso", parent=estilos["Normal"], fontName="Helvetica-Oblique",
        fontSize=8, textColor=cor_texto_secundario, spaceBefore=5,
    )
    estilo_vazio = ParagraphStyle(
        "SemDados", parent=estilos["Normal"], fontName="Helvetica-Oblique",
        fontSize=9.5, textColor=cor_texto_secundario,
    )
    estilo_descricao_secao = ParagraphStyle(
        "DescricaoSecao", parent=estilos["Normal"], fontName="Helvetica-Oblique",
        fontSize=9, textColor=cor_texto_secundario, spaceBefore=2, spaceAfter=6,
    )
    # Células de tabela: Paragraph (não string crua) para quebrar linha —
    # sem isso, cabeçalhos longos ("Períodos Consecutivos em Queda") não
    # cabem na largura estreita de uma coluna numérica e vazam por cima do
    # cabeçalho vizinho, em vez de quebrar em 2-3 linhas.
    estilo_celula_texto = ParagraphStyle(
        "CelulaTexto", parent=estilos["Normal"], fontName="Helvetica",
        fontSize=8.5, leading=10.5, alignment=TA_LEFT,
    )
    estilo_celula_numero = ParagraphStyle(
        "CelulaNumero", parent=estilos["Normal"], fontName="Helvetica",
        fontSize=8.5, leading=10.5, alignment=TA_RIGHT,
    )
    estilo_cabecalho_texto = ParagraphStyle(
        "CabecalhoTexto", parent=estilos["Normal"], fontName="Helvetica-Bold",
        fontSize=8.5, leading=10.5, alignment=TA_LEFT, textColor=colors.white,
    )
    estilo_cabecalho_numero = ParagraphStyle(
        "CabecalhoNumero", parent=estilos["Normal"], fontName="Helvetica-Bold",
        fontSize=8.5, leading=10.5, alignment=TA_RIGHT, textColor=colors.white,
    )

    elementos = []

    # ------------------------------------------------------------------
    # Capa
    # ------------------------------------------------------------------
    agora = datetime.now()
    elementos.append(Spacer(1, 1.8 * cm))
    if os.path.exists(CAMINHO_LOGO):
        logo_capa = Image(CAMINHO_LOGO, width=3.4 * cm, height=3.4 * cm * 306 / 572)
        logo_capa.hAlign = "CENTER"
        elementos.append(logo_capa)
        elementos.append(Spacer(1, 20))

    # Faixa preta com o nome do sistema — mesma cor de destaque usada nos
    # cabeçalhos das tabelas, para a capa e o conteúdo lerem como um só
    # documento visualmente coerente.
    estilo_titulo_faixa = ParagraphStyle(
        "TituloFaixa", parent=estilos["Title"], fontName="Helvetica-Bold",
        fontSize=26, leading=30, textColor=colors.white, alignment=TA_CENTER, spaceAfter=4,
    )
    estilo_subtitulo_faixa = ParagraphStyle(
        "SubtituloFaixa", parent=estilos["Normal"], fontName="Helvetica",
        fontSize=12, leading=15, textColor=colors.HexColor("#CCCCCC"), alignment=TA_CENTER,
    )
    faixa_titulo = Table(
        [[[Paragraph(NOME_SISTEMA, estilo_titulo_faixa), Paragraph("Relatório Padrão", estilo_subtitulo_faixa)]]],
        colWidths=[largura_util],
    )
    faixa_titulo.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.black),
        ("TOPPADDING", (0, 0), (-1, -1), 16),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 16),
    ]))
    elementos.append(faixa_titulo)
    elementos.append(Spacer(1, 26))

    if nome_empresa:
        elementos.append(Paragraph(nome_empresa, estilo_empresa_capa))
        elementos.append(Spacer(1, 16))
    elementos.append(HRFlowable(width=5 * cm, thickness=1.4, color=cor_marca, hAlign="CENTER"))
    elementos.append(Spacer(1, 18))

    # Cartão de metadados (quem gerou, quando) — visualmente separado do
    # título/empresa, para ficar claro que é informação de rodapé da capa.
    colunas_metadados = []
    if nome_usuario:
        colunas_metadados.append([Paragraph("Gerado por", estilo_rotulo_metadados), Paragraph(nome_usuario, estilo_valor_metadados)])
    colunas_metadados.append([
        Paragraph("Data e hora de extração", estilo_rotulo_metadados),
        Paragraph(agora.strftime("%d/%m/%Y às %H:%M"), estilo_valor_metadados),
    ])
    largura_coluna_metadados = largura_util / (2 * len(colunas_metadados))
    cartao_metadados = Table(
        [colunas_metadados], colWidths=[largura_coluna_metadados * 2] * len(colunas_metadados), hAlign="CENTER",
    )
    cartao_metadados.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
    ]))
    elementos.append(cartao_metadados)
    elementos.append(Spacer(1, 14))
    elementos.append(Paragraph(NOME_EMPRESA, estilo_meta_capa))
    elementos.append(PageBreak())

    # ------------------------------------------------------------------
    # Índice — só compensa a partir de uns 3-4 relatórios; com 1 ou 2, a
    # página extra vira um clique a mais sem ganho real de navegação.
    # ------------------------------------------------------------------
    total_secoes = sum(len(analises) for analises in resultados_analise.values())
    if total_secoes > 2:
        estilo_titulo_indice = ParagraphStyle(
            "TituloIndice", parent=estilos["Heading1"], fontName="Helvetica-Bold",
            fontSize=16, textColor=cor_marca, spaceAfter=14,
        )
        estilo_entrada_indice = ParagraphStyle(
            "EntradaIndice", parent=estilos["Normal"], fontName="Helvetica",
            fontSize=10.5, leading=16, textColor=colors.HexColor("#1a1a1a"),
        )
        indice = TableOfContents()
        indice.levelStyles = [estilo_entrada_indice]
        # TableOfContents desenha "texto ..... nº" com um separador de pontos
        # próprio — não precisa de tabela/grid manual.
        elementos.append(Paragraph("Índice", estilo_titulo_indice))
        elementos.append(indice)
        elementos.append(PageBreak())

    # ------------------------------------------------------------------
    # Seções (uma por análise x granularidade)
    # ------------------------------------------------------------------
    for granularidade, analises in resultados_analise.items():
        for chave, df in analises.items():
            titulo = nomes_analise.get(chave, chave).replace("_", " ")
            descricao = descricao_analise.get(chave)

            paragrafo_titulo_secao = Paragraph(titulo, estilo_secao)
            paragrafo_titulo_secao._indice_texto = f"{titulo} — {granularidade}"
            itens_cabecalho = [paragrafo_titulo_secao, Paragraph(granularidade, estilo_granularidade)]
            if descricao:
                itens_cabecalho.append(Paragraph(descricao, estilo_descricao_secao))
            itens_cabecalho.append(Spacer(1, 6))
            itens_cabecalho.append(HRFlowable(width="100%", thickness=1.2, color=cor_marca, spaceAfter=10))
            cabecalho_secao = KeepTogether(itens_cabecalho)
            elementos.append(cabecalho_secao)

            if df is None or df.empty:
                elementos.append(Paragraph("Sem dados para esta análise/granularidade.", estilo_vazio))
                elementos.append(Spacer(1, 22))
                continue

            df_limitado, total = _limitar(df)
            colunas_moeda = set(colunas_moeda_por_analise.get(chave, []))
            colunas = list(df_limitado.columns)
            eh_numerica = [
                (coluna in colunas_moeda) or pd.api.types.is_numeric_dtype(df_limitado[coluna])
                for coluna in colunas
            ]

            titulos_colunas = [_titulo_coluna(coluna) for coluna in colunas]
            linha_cabecalho = [
                Paragraph(titulo, estilo_cabecalho_numero if numerica else estilo_cabecalho_texto)
                for titulo, numerica in zip(titulos_colunas, eh_numerica)
            ]
            textos_por_coluna = {coluna: [] for coluna in colunas}
            linhas_formatadas = []
            for _, linha in df_limitado.iterrows():
                linha_fmt = []
                for coluna, numerica in zip(colunas, eh_numerica):
                    valor = linha[coluna]
                    if isinstance(valor, bool):
                        texto_valor = "Sim" if valor else "Não"
                    elif coluna in colunas_moeda:
                        texto_valor = _formatar_moeda_br(valor)
                    elif numerica and isinstance(valor, float):
                        texto_valor = _formatar_numero_br(valor)
                    else:
                        texto_valor = str(valor)
                    textos_por_coluna[coluna].append(texto_valor)
                    estilo_celula = estilo_celula_numero if numerica else estilo_celula_texto
                    linha_fmt.append(Paragraph(texto_valor, estilo_celula))
                linhas_formatadas.append(linha_fmt)
            dados_tabela = [linha_cabecalho] + linhas_formatadas

            # Largura por coluna baseada no CONTEÚDO real, não num split
            # uniforme entre colunas de texto — antes, "Grupo" (valores
            # curtos como "Demais") recebia a mesma largura que "Cliente"
            # (razões sociais longas), sobrando espaço em branco enorme à
            # direita da tabela em vez de ir pra quem precisa. Numéricas
            # usam o valor MÁXIMO formatado da própria coluna (não um
            # tamanho fixo de 2,3cm) — é isso que também corrige colunas de
            # moeda quebrando em 2 linhas pra valores grandes e ficando numa
            # linha só pra valores pequenos, dentro da mesma tabela.
            padding_celula = 14 + 6  # LEFTPADDING + RIGHTPADDING da tabela + margem de segurança, em pontos
            larguras_minimas = [_largura_maior_palavra(titulo) + padding_celula for titulo in titulos_colunas]
            LARGURA_MIN_COLUNA = 2.3 * cm
            LARGURA_MAX_COLUNA_TEXTO = 9.0 * cm
            larguras_base = []
            for coluna, titulo_largura_minima, numerica in zip(colunas, larguras_minimas, eh_numerica):
                textos = textos_por_coluna[coluna]
                if numerica:
                    largura_conteudo = max((stringWidth(t, "Helvetica", 8.5) for t in textos), default=0)
                else:
                    largura_conteudo = _largura_tipica_conteudo(textos)
                largura = max(titulo_largura_minima, largura_conteudo + padding_celula, LARGURA_MIN_COLUNA)
                if not numerica:
                    largura = min(largura, LARGURA_MAX_COLUNA_TEXTO)
                larguras_base.append(largura)

            soma_base = sum(larguras_base)
            if soma_base > largura_util:
                # Corta só a "sobra" de cada coluna (base − mínimo da maior
                # palavra do cabeçalho) — nunca abaixo do mínimo, senão volta
                # a quebrar palavra no meio (era o que um corte uniforme
                # fazia antes, desfazendo o próprio mínimo calculado acima).
                excesso = soma_base - largura_util
                sobras = [b - m for b, m in zip(larguras_base, larguras_minimas)]
                sobra_total = sum(sobras)
                if sobra_total > 0:
                    fator_corte = min(excesso / sobra_total, 1.0)
                    larguras = [b - s * fator_corte for b, s in zip(larguras_base, sobras)]
                else:
                    larguras = larguras_base
            elif soma_base < largura_util:
                # Sobra espaço: distribui só entre as colunas de TEXTO,
                # proporcional à largura-base de cada uma — não entre as
                # numéricas, que já têm exatamente a largura que seu maior
                # valor formatado precisa (dar mais só abriria vazio dentro
                # da própria célula, sem ganho nenhum de legibilidade).
                indices_texto = [indice for indice, numerica in enumerate(eh_numerica) if not numerica]
                if indices_texto:
                    sobra = largura_util - soma_base
                    pesos = [larguras_base[indice] for indice in indices_texto]
                    soma_pesos = sum(pesos) or 1
                    larguras = list(larguras_base)
                    for indice, peso in zip(indices_texto, pesos):
                        larguras[indice] += sobra * (peso / soma_pesos)
                else:
                    larguras = larguras_base
            else:
                larguras = larguras_base

            tabela = Table(dados_tabela, colWidths=larguras, repeatRows=1)
            estilo_tabela = [
                ("BACKGROUND", (0, 0), (-1, 0), cor_cabecalho_tabela),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("LINEBELOW", (0, 0), (-1, 0), 1.2, cor_cabecalho_tabela),
                ("LINEBELOW", (0, 1), (-1, -2), 0.4, cor_linha_sutil),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, cor_linha_alternada]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
            for indice, numerica in enumerate(eh_numerica):
                estilo_tabela.append(("ALIGN", (indice, 0), (indice, -1), "RIGHT" if numerica else "LEFT"))
            tabela.setStyle(TableStyle(estilo_tabela))
            elementos.append(tabela)

            if total > MAX_LINHAS_TABELA:
                elementos.append(Paragraph(
                    f"Mostrando {MAX_LINHAS_TABELA} de {total} registros — para a base completa, exporte em Excel.",
                    estilo_aviso,
                ))
            elementos.append(Spacer(1, 24))

    doc.multiBuild(elementos, onFirstPage=lambda c, d: None, onLaterPages=_cabecalho_rodape)


def exportar_relatorio_word(caminho_saida, resultados_analise, nomes_analise, nome_usuario="", colunas_moeda_por_analise=None, nome_empresa="", descricao_analise=None):
    import docx
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import pandas as pd

    colunas_moeda_por_analise = colunas_moeda_por_analise or {}
    descricao_analise = descricao_analise or {}
    cor_marca = RGBColor(0x1F, 0x4E, 0x78)
    cor_texto_secundario = RGBColor(0x6B, 0x72, 0x80)

    documento = docx.Document()

    if os.path.exists(CAMINHO_LOGO):
        paragrafo_logo = documento.add_paragraph()
        paragrafo_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragrafo_logo.add_run().add_picture(CAMINHO_LOGO, width=Cm(3.2))

    titulo = documento.add_heading(NOME_SISTEMA, level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.color.rgb = cor_marca

    subtitulo = documento.add_paragraph("Relatório Padrão")
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.runs[0].font.color.rgb = cor_texto_secundario

    if nome_empresa:
        p_empresa = documento.add_paragraph(nome_empresa)
        p_empresa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_empresa.runs[0].font.color.rgb = cor_marca
        p_empresa.runs[0].font.bold = True
        p_empresa.runs[0].font.size = Pt(14)

    if nome_usuario:
        p = documento.add_paragraph(f"Gerado por {nome_usuario}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.color.rgb = cor_texto_secundario
    p = documento.add_paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')} — {NOME_EMPRESA}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = cor_texto_secundario
    documento.add_page_break()

    for granularidade, analises in resultados_analise.items():
        for chave, df in analises.items():
            titulo_analise = nomes_analise.get(chave, chave).replace("_", " ")
            cabecalho = documento.add_heading(titulo_analise, level=2)
            for run in cabecalho.runs:
                run.font.color.rgb = cor_marca
            subtitulo_secao = documento.add_paragraph(granularidade)
            subtitulo_secao.runs[0].font.color.rgb = cor_texto_secundario
            subtitulo_secao.runs[0].font.size = Pt(9.5)

            descricao = descricao_analise.get(chave)
            if descricao:
                p_descricao = documento.add_paragraph(descricao)
                p_descricao.runs[0].italic = True
                p_descricao.runs[0].font.color.rgb = cor_texto_secundario
                p_descricao.runs[0].font.size = Pt(9)

            if df is None or df.empty:
                p = documento.add_paragraph("Sem dados para esta análise/granularidade.")
                p.runs[0].italic = True
                continue

            df_limitado, total = _limitar(df)
            colunas_moeda = set(colunas_moeda_por_analise.get(chave, []))
            colunas = list(df_limitado.columns)
            eh_numerica = [
                (coluna in colunas_moeda) or pd.api.types.is_numeric_dtype(df_limitado[coluna])
                for coluna in colunas
            ]

            tabela = documento.add_table(rows=1, cols=len(colunas))
            try:
                tabela.style = "Light Grid Accent 1"
            except KeyError:
                pass
            celulas_cabecalho = tabela.rows[0].cells
            for i, coluna in enumerate(colunas):
                celulas_cabecalho[i].text = _titulo_coluna(coluna)
                celulas_cabecalho[i].paragraphs[0].runs[0].bold = True

            for _, linha in df_limitado.iterrows():
                celulas = tabela.add_row().cells
                for i, coluna in enumerate(colunas):
                    valor = linha[coluna]
                    if isinstance(valor, bool):
                        texto_valor = "Sim" if valor else "Não"
                    elif coluna in colunas_moeda:
                        texto_valor = _formatar_moeda_br(valor)
                    elif eh_numerica[i] and isinstance(valor, float):
                        texto_valor = _formatar_numero_br(valor)
                    else:
                        texto_valor = str(valor)
                    celulas[i].text = texto_valor
                    if eh_numerica[i]:
                        celulas[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if total > MAX_LINHAS_TABELA:
                paragrafo = documento.add_paragraph(
                    f"Mostrando {MAX_LINHAS_TABELA} de {total} registros — para a base completa, exporte em Excel."
                )
                paragrafo.runs[0].italic = True
            documento.add_paragraph()

    documento.save(caminho_saida)
